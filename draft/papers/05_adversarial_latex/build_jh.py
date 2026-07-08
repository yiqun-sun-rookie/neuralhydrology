"""Build a clean Journal-of-Hydrology (Elsevier "Your Paper Your Way") Word manuscript
from main.tex. Not the AGU template: plain, complete, submittable .docx.

pandoc (math -> OMML, citeproc from references.bib) + python-docx front matter + 3-line tables.
Run:  python build_jh.py   ->  _build/manuscript_JH.docx
"""
import re
import subprocess
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).resolve().parent
BUILD = HERE / "_build"
(BUILD / "figures").mkdir(parents=True, exist_ok=True)

TITLE = ("Observation-Scale Forcing Errors Can Substantially Degrade LSTM "
         "Rainfall-Runoff Predictions: A Multi-Level Adversarial Stress Test")
AUTHOR = "Y. Sun"
AFFIL = "College of Hydrology and Water Resources, Hohai University, Nanjing, China"
CORRESP = "Corresponding author: Y. Sun (sunyiq@hhu.edu.cn)"
KEYWORDS = ("Rainfall-runoff modeling; Long short-term memory; Adversarial robustness; "
            "Input uncertainty; Quality control; Streamflow prediction")
HIGHLIGHTS = [
    "Gradient-based stress testing finds about 16 times the loss of matched random noise",
    "Mean- and variance-preserving input errors evade those checks yet keep 78% of damage",
    "Failures concentrate on temperature inputs and snowmelt flood peaks in snow basins",
]

tex = (HERE / "main.tex").read_text(encoding="utf-8")
aux = (HERE / "main.aux").read_text(encoding="utf-8")

# ---- 1. cross-reference numbers from the compiled .aux (authoritative) ----
REF = {m.group(1): m.group(2) for m in re.finditer(r"\\newlabel\{([^}]+)\}\{\{([^}]*)\}", aux)}

# ---- 2. figures: PDF -> 300 dpi PNG ----
for pdf in sorted((HERE / "figures").glob("fig*.pdf")):
    d = fitz.open(pdf)
    d[0].get_pixmap(dpi=300).save(BUILD / "figures" / f"{pdf.stem}.png")
    d.close()

# ---- 3. abstract + body ----
abstract = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", tex, re.S).group(1).strip()
body = tex[tex.index("\\section{Introduction}"):tex.index("\\end{document}")]


def conv_cites(t):
    # pandoc's latex reader handles \cite/\citet/\citep natively via --citeproc;
    # only apacite's \citeA (textual) and \citeNP need mapping to \citet.
    t = t.replace("\\citeA{", "\\citet{").replace("\\citeNP{", "\\citet{")
    return t


def preprocess(t):
    t = conv_cites(t)
    t = re.sub(r"\\ref\{([^}]+)\}", lambda m: REF.get(m.group(1), "??"), t)
    t = re.sub(r"\\includegraphics\[[^\]]*\]\{figures/([^}]+)\}",
               lambda m: "\\includegraphics{" + (BUILD / "figures" / (Path(m.group(1)).stem + ".png")).as_posix() + "}", t)
    t = re.sub(r"\\label\{[^}]+\}", "", t)
    t = t.replace("\\acknowledgments", "\\section*{Acknowledgments}")
    t = re.sub(r"\\bibliography\{[^}]+\}", lambda m: "\\section*{References}", t)
    return t


hl_tex = "\\section*{Highlights}\n\\begin{itemize}\n" + "\n".join(f"\\item {h}" for h in HIGHLIGHTS) + "\n\\end{itemize}\n"
front = (f"\\section*{{Abstract}}\n{conv_cites(abstract)}\n\n"
         f"\\noindent\\textbf{{Keywords:}} {KEYWORDS}\n\n" + hl_tex + "\n")
pre = front + preprocess(body)
(BUILD / "pre.tex").write_text(pre, encoding="utf-8")

# ---- 4. pandoc ----
out = BUILD / "body.docx"
subprocess.run(["pandoc", str(BUILD / "pre.tex"), "-o", str(out),
                "-f", "latex", "-t", "docx", "--citeproc",
                "--bibliography", str(HERE / "references.bib"),
                "--number-sections"], check=True, cwd=str(HERE))

# ---- 5. python-docx: front matter + tables + captions ----
doc = Document(str(out))
first = doc.paragraphs[0]


def ins(text, before, bold=False, italic=False, size=None, align=None, color=None):
    p = before.insert_paragraph_before()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    if align is not None:
        p.alignment = align
    return p


# insert_paragraph_before(first) always lands immediately before `first`,
# so call in VISUAL order: Title first ends up on top.
ins(TITLE, first, bold=True, size=16)
ins(AUTHOR, first, size=12)
ins(AFFIL, first, size=10)
ins(CORRESP, first, italic=True, size=10)
ins("", first)  # spacer before Abstract

# 3-line tables + centered
def set_border(cell, edge, sz):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{e}"))
        if el is None:
            el = OxmlElement(f"w:{e}")
            borders.append(el)
        if e == edge:
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz)); el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")


for tbl in doc.tables:
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = tbl.rows
    for c in rows[0].cells:
        set_border(c, "top", 12)
        set_border(c, "bottom", 4)
    for c in rows[-1].cells:
        set_border(c, "bottom", 12)

# bold "Figure N."/"Table N." caption prefixes (pandoc puts captions as Image/Table caption styles)
for p in doc.paragraphs:
    sty = (p.style.name or "").lower()
    if "caption" in sty and p.runs:
        p.runs[0].bold = True

doc.save(str(BUILD / "manuscript_JH.docx"))
print("WROTE", BUILD / "manuscript_JH.docx")
print("tables:", len(doc.tables), "| figures(drawings):",
      len(doc.element.body.findall('.//' + qn('w:drawing'))))
