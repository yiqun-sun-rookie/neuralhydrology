"""Build an AGU-Word-template .docx from main.tex, with the running header set to
"Journal of Hydrology" (=> "manuscript submitted to Journal of Hydrology").

Reuses the pandoc + citeproc mechanism (references rebuilt from references.bib,
inline math -> native Word OMML) and adds the AGU template styling + front matter
per the latex-to-agu-word skill (style remap, Key Points bullets via numId 9,
3-line tables, header journal-name fix).

Run:  python build_agu.py   ->  manuscript_AGU.docx
"""
import re
import subprocess
import zipfile
import shutil
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
BUILD = HERE / "_build"
(BUILD / "figures").mkdir(parents=True, exist_ok=True)
TEMPLATE = Path("G:/github/pycharm/projects/nature_1st/draft/wrr_agu_template_2026-05-12/AGUWordTemplate.docx")
OUT = HERE / "manuscript_AGU.docx"
JOURNAL = "Journal of Hydrology"

TITLE = ("Observation-Scale Forcing Errors Can Substantially Degrade LSTM "
         "Rainfall-Runoff Predictions: A Multi-Level Adversarial Stress Test")
AUTHORS = [("Y. Sun", False), ("1", True)]
AFFILIATIONS = [("1", "College of Hydrology and Water Resources, Hohai University, Nanjing, China.")]
CORRESPONDING = "Corresponding author: Y. Sun (sunyiq@hhu.edu.cn)"
KEY_POINTS = [
    "Gradient-based stress testing reveals 16 times the worst-case NSE loss of matched random noise for a 531-basin CAMELS-US LSTM at eps=0.1",
    "Perturbations preserving each input's mean and standard deviation evade checks based on those statistics yet retain 78% of the damage",
    "The failure modes are hydrologically interpretable, concentrating on the temperature channel and melt-driven flood peaks in snow catchments",
]

tex = (HERE / "main.tex").read_text(encoding="utf-8")
aux = (HERE / "main.aux").read_text(encoding="utf-8")
REF = {m.group(1): m.group(2) for m in re.finditer(r"\\newlabel\{([^}]+)\}\{\{([^}]*)\}", aux)}

# 1. figures -> PNG
for pdf in sorted((HERE / "figures").glob("fig*.pdf")):
    d = fitz.open(pdf)
    d[0].get_pixmap(dpi=300).save(BUILD / "figures" / f"{pdf.stem}.png")
    d.close()

# 2. abstract + PLS + body
abstract = re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", tex, re.S).group(1).strip()
pls = re.search(r"\\section\*\{Plain Language Summary\}(.*?)(?=\\section|\%\%|$)", tex, re.S)
pls_txt = pls.group(1).strip() if pls else ""
body = tex[tex.index("\\section{Introduction}"):tex.index("\\end{document}")]


def conv_cites(t):
    return t.replace("\\citeA{", "\\citet{").replace("\\citeNP{", "\\citet{")


def preprocess(t):
    t = conv_cites(t)
    t = re.sub(r"\\ref\{([^}]+)\}", lambda m: REF.get(m.group(1), "??"), t)
    t = re.sub(r"\\includegraphics\[[^\]]*\]\{figures/([^}]+)\}",
               lambda m: "\\includegraphics{" + (BUILD / "figures" / (Path(m.group(1)).stem + ".png")).as_posix() + "}", t)
    t = re.sub(r"\\label\{[^}]+\}", "", t)
    # drop the AGU PLS section from body (rebuilt separately); drop leftover
    t = re.sub(r"\\section\*\{Plain Language Summary\}.*?(?=\\section)", "", t, flags=re.S)
    t = t.replace("\\acknowledgments", "\\section*{Acknowledgments}")
    t = re.sub(r"\\bibliography\{[^}]+\}", lambda m: "\\section*{References}", t)
    return t


pre = f"\\section*{{Abstract}}\n{conv_cites(abstract)}\n\n"
if pls_txt:
    pre += f"\\section*{{Plain Language Summary}}\n{conv_cites(pls_txt)}\n\n"
pre += preprocess(body)
(BUILD / "pre.tex").write_text(pre, encoding="utf-8")
(BUILD / "nocite.yaml").write_text("nocite: '@*'\n", encoding="utf-8")

# 3. pandoc with AGU reference-doc
inter = BUILD / "intermediate.docx"
subprocess.run(["pandoc", str(BUILD / "pre.tex"), "-o", str(inter),
                "-f", "latex", "-t", "docx", "--reference-doc", str(TEMPLATE),
                "--citeproc", "--bibliography", str(HERE / "references.bib"),
                "--metadata-file", str(BUILD / "nocite.yaml"), "--number-sections"],
               check=True, cwd=str(HERE))

# 4. python-docx post-process
doc = Document(str(inter))

# pandoc with the AGU reference-doc flattens everything to "Normal" (the template
# lacks pandoc's expected Heading styleIds), so rebuild structure from the
# --number-sections heading text ("1  Intro", "2.1  ...", "2.2.1  ...") and from
# caption / body cues.
HEAD_RE = re.compile(r"^\d+(\.\d+)*[\t ]+\S")


def restyle(name):
    return doc.styles[name] if name in [s.name for s in doc.styles] else None


AGU_TEXT = restyle("Text")
AGU_MAIN = restyle("Heading-Main")
AGU_SEC = restyle("Heading-Secondary")
AGU_CAP = restyle("FigureorTableCaption")
for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    sid = p.style.name
    if HEAD_RE.match(t):
        lvl = t.split()[0].count(".") + 1
        p.style = AGU_MAIN if lvl == 1 else AGU_SEC
    elif t in ("References", "Acknowledgments", "Open Research"):
        if AGU_MAIN:
            p.style = AGU_MAIN
    elif "caption" in sid.lower() or t[:7] in ("Figure ", "Table 1", "Table 2", "Table 3", "Table 4", "Table 5"):
        if AGU_CAP:
            p.style = AGU_CAP
    elif sid == "Normal":
        if AGU_TEXT:
            p.style = AGU_TEXT

# The abstract heading "Abstract" comes out of pandoc as a body paragraph.
# We PREPEND the front matter (Title..Key Points) before that first "Abstract"
# paragraph, in strict visual order, each new paragraph inserted before the SAME
# fixed anchor so appended order == visual order.
anchor = next(p for p in doc.paragraphs if p.text.strip() == "Abstract")


def style_of(name):
    try:
        return doc.styles[name]
    except KeyError:
        return None


def add(text, style=None, before=anchor):
    p = before.insert_paragraph_before(text)
    s = style_of(style) if style else None
    if s is not None:
        p.style = s
    return p


# Title
add(TITLE, "Title")
# Authors (with superscripts)
pa = anchor.insert_paragraph_before()
if style_of("Authors"):
    pa.style = style_of("Authors")
for text, sup in AUTHORS:
    r = pa.add_run(text)
    if sup:
        r.font.superscript = True
# Affiliations + corresponding
for num, aff in AFFILIATIONS:
    p = anchor.insert_paragraph_before()
    if style_of("Affiliation"):
        p.style = style_of("Affiliation")
    r = p.add_run(num); r.font.superscript = True
    p.add_run(aff)
add(CORRESPONDING, "Affiliation")
# Key Points
add("Key Points:", "Heading-Main")
for kp in KEY_POINTS:
    p = anchor.insert_paragraph_before(kp)
    if style_of("Key Points"):
        p.style = style_of("Key Points")
    numPr = OxmlElement("w:numPr")
    il = OxmlElement("w:ilvl"); il.set(qn("w:val"), "0"); numPr.append(il)
    ni = OxmlElement("w:numId"); ni.set(qn("w:val"), "9"); numPr.append(ni)
    p._p.get_or_add_pPr().append(numPr)

# Style the Abstract / Plain Language Summary headings + bodies (pandoc emitted them
# as body paragraphs from the \section* in pre.tex). Heading -> Heading-Main,
# following body paragraph -> Abstract style.
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() in ("Abstract", "Plain Language Summary"):
        if style_of("Heading-Main"):
            p.style = style_of("Heading-Main")
        for q in doc.paragraphs[i + 1:]:
            if q.text.strip():
                if style_of("Abstract"):
                    q.style = style_of("Abstract")
                break

# 3-line tables
def set_border(cell, edge, sz):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = OxmlElement("w:tcBorders"); tcPr.append(b)
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = b.find(qn(f"w:{e}"))
        if el is None:
            el = OxmlElement(f"w:{e}"); b.append(el)
        if e == edge:
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), str(sz)); el.set(qn("w:space"), "0"); el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")


for tbl in doc.tables:
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c in tbl.rows[0].cells:
        set_border(c, "top", 12); set_border(c, "bottom", 4)
    for c in tbl.rows[-1].cells:
        set_border(c, "bottom", 12)

# bold Figure/Table caption prefixes
for p in doc.paragraphs:
    if "caption" in (p.style.name or "").lower() and p.runs:
        p.runs[0].bold = True

# references -> Reference style
ref_i = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip().lower() == "references"), None)
if ref_i is not None:
    for p in doc.paragraphs[ref_i + 1:]:
        if p.text.strip():
            try:
                p.style = doc.styles["Reference"]
            except KeyError:
                pass

doc.save(str(OUT))

# 5. header journal name (zip-level string replace)
tmp = BUILD / "hdr.docx"
shutil.copy(OUT, tmp)
zin = zipfile.ZipFile(tmp)
zout = zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED)
for item in zin.namelist():
    data = zin.read(item)
    if item.startswith("word/header") and item.endswith(".xml"):
        txt = data.decode("utf-8")
        # placeholder is split across several <w:t> runs ("replace this text with
        # name of" / "AGU journal"); collapse the run sequence to the journal name.
        txt = re.sub(r"(<w:t[^>]*>)replace this text with name of\s*(</w:t>)",
                     r"\g<1>" + JOURNAL + r"\g<2>", txt)
        txt = re.sub(r"<w:t[^>]*>\s*AGU journal\s*</w:t>", "<w:t></w:t>", txt)
        data = txt.encode("utf-8")
    zout.writestr(item, data)
zin.close(); zout.close(); tmp.unlink()

print("WROTE", OUT)
print("tables:", len(doc.tables), "| figures:",
      len(doc.element.body.findall('.//' + qn('w:drawing'))))
